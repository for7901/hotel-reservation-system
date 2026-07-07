# -*- coding: utf-8 -*-
"""生成全栈项目主题讨论会议记录 Word 模板"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\17879\Desktop\会议记录.docx"

# 统一格式常量
FONT_NAME = "宋体"
FONT_SIZE_TITLE = Pt(16)      # 文档标题
FONT_SIZE_HEADING = Pt(14)    # 一级章节
FONT_SIZE_SUB = Pt(12)      # 二级/正文
FONT_SIZE_TABLE = Pt(11)    # 表格内容
FONT_SIZE_NOTE = Pt(10)     # 备注说明


def set_run_font(run, size, bold=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, FONT_SIZE_TITLE, bold=True)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, FONT_SIZE_HEADING, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, FONT_SIZE_SUB, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, FONT_SIZE_SUB)
    if indent:
        p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_blank(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)


def set_cell_text(cell, text, bold=False, align_center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, FONT_SIZE_TABLE, bold=bold)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def add_info_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], label, bold=True)
        set_cell_text(table.rows[i].cells[1], value)
        table.rows[i].cells[0].width = Cm(3.5)
        table.rows[i].cells[1].width = Cm(12)
    add_blank(doc)
    return table


def add_agenda_table(doc):
    headers = ["序号", "议程项", "汇报/主讲人", "预计时长"]
    data = [
        ["1", "项目背景与目标说明", "", "10 min"],
        ["2", "全栈项目主题头脑风暴", "", "30 min"],
        ["3", "候选主题评估与筛选", "", "20 min"],
        ["4", "技术栈初步讨论", "", "15 min"],
        ["5", "分工与下一步计划", "", "10 min"],
    ]
    table = doc.add_table(rows=1 + len(data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, align_center=True)
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], val, align_center=(j == 0))
    add_blank(doc)
    return table


def add_action_table(doc):
    headers = ["序号", "待办内容", "负责人", "截止日期", "优先级", "状态"]
    data = [["1", "", "", "YYYY-MM-DD", "高/中/低", "待开始"]] * 4
    table = doc.add_table(rows=1 + len(data), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, align_center=True)
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], val, align_center=(j in (0, 3, 4, 5)))
    add_blank(doc)
    return table


def add_risk_table(doc):
    headers = ["序号", "问题/风险描述", "影响程度", "应对措施", "跟进人"]
    data = [["1", "", "高/中/低", "", ""]] * 2
    table = doc.add_table(rows=1 + len(data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, align_center=True)
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], val, align_center=(j in (0, 2)))
    add_blank(doc)
    return table


def add_idea_table(doc):
    headers = ["序号", "项目主题名称", "核心功能描述", "目标用户", "技术亮点", "可行性", "投票/评分"]
    data = [["1", "", "", "", "", "高/中/低", ""]] * 5
    table = doc.add_table(rows=1 + len(data), cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, align_center=True)
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], val, align_center=(j in (0, 5, 6)))
    add_blank(doc)
    return table


def add_tech_table(doc):
    headers = ["层次", "候选技术", "选定方案", "备注"]
    data = [
        ["前端", "Vue / React / Angular 等", "", ""],
        ["后端", "Spring Boot / Node.js / Django 等", "", ""],
        ["数据库", "MySQL / PostgreSQL / MongoDB 等", "", ""],
        ["缓存/中间件", "Redis / RabbitMQ 等", "", ""],
        ["部署", "Docker / 云服务器 等", "", ""],
    ]
    table = doc.add_table(rows=1 + len(data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, align_center=True)
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], val)
    add_blank(doc)
    return table


def main():
    doc = Document()
    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 默认正文样式
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = FONT_SIZE_SUB

    add_title(doc, "全栈项目主题讨论 — 会议记录")

    add_heading(doc, "一、会议基本信息")
    add_info_table(doc, [
        ("会议主题", "全栈项目主题头脑风暴与选型讨论"),
        ("会议类型", "□ 头脑风暴  □ 项目立项  □ 技术评审  □ 其他：________"),
        ("会议时间", "YYYY-MM-DD  HH:MM — HH:MM"),
        ("会议地点/方式", "□ 线下（地点：________）  □ 线上（平台：________）"),
        ("主持人", ""),
        ("记录人", ""),
        ("参会人员", ""),
        ("缺席人员", ""),
    ])

    add_heading(doc, "二、会议议程")
    add_agenda_table(doc)

    add_heading(doc, "三、项目背景与目标")
    add_body(doc, "【背景说明】（简述为何需要确定全栈项目主题，如课程设计、毕业设计、团队练手等）")
    add_body(doc, "")
    add_body(doc, "【项目目标】")
    add_body(doc, "1. 确定一个具有实际应用价值的全栈项目主题；", indent=True)
    add_body(doc, "2. 明确目标用户群体与核心功能范围；", indent=True)
    add_body(doc, "3. 初步评估技术可行性与开发周期；", indent=True)
    add_body(doc, "4. 为后续需求分析与技术选型奠定基础。", indent=True)
    add_blank(doc)

    add_heading(doc, "四、项目主题头脑风暴")
    add_sub_heading(doc, "4.1 讨论规则")
    add_body(doc, "• 鼓励发散思维，先数量后质量，暂不否定任何想法；", indent=True)
    add_body(doc, "• 每个主题需说明：解决什么问题、面向谁、核心功能是什么；", indent=True)
    add_body(doc, "• 记录所有提出的候选主题，后续统一评估筛选。", indent=True)
    add_blank(doc)

    add_sub_heading(doc, "4.2 候选主题汇总")
    add_idea_table(doc)

    add_sub_heading(doc, "4.3 主题讨论要点")
    add_body(doc, "【主题一：________】")
    add_body(doc, "讨论要点：", indent=True)
    add_body(doc, "• ", indent=True)
    add_body(doc, "• ", indent=True)
    add_body(doc, "优势：", indent=True)
    add_body(doc, "• ", indent=True)
    add_body(doc, "顾虑/风险：", indent=True)
    add_body(doc, "• ", indent=True)
    add_blank(doc)

    add_body(doc, "【主题二：________】")
    add_body(doc, "讨论要点：", indent=True)
    add_body(doc, "• ", indent=True)
    add_body(doc, "优势：", indent=True)
    add_body(doc, "• ", indent=True)
    add_body(doc, "顾虑/风险：", indent=True)
    add_body(doc, "• ", indent=True)
    add_blank(doc)

    add_heading(doc, "五、主题评估与决议")
    add_sub_heading(doc, "5.1 评估维度")
    add_body(doc, "• 业务价值：是否解决真实问题，有无实际应用场景；", indent=True)
    add_body(doc, "• 技术覆盖：能否涵盖前端、后端、数据库等全栈要素；", indent=True)
    add_body(doc, "• 开发难度：团队技能是否匹配，周期是否合理；", indent=True)
    add_body(doc, "• 可扩展性：后续是否有迭代和功能扩展空间；", indent=True)
    add_body(doc, "• 差异化：与常见项目是否有一定特色。", indent=True)
    add_blank(doc)

    add_sub_heading(doc, "5.2 最终决议")
    add_info_table(doc, [
        ("选定项目主题", ""),
        ("项目一句话描述", ""),
        ("目标用户群体", ""),
        ("核心功能范围（MVP）", ""),
        ("预计开发周期", ""),
        ("未采纳主题及原因", ""),
    ])

    add_heading(doc, "六、技术栈初步讨论")
    add_tech_table(doc)
    add_body(doc, "【技术选型说明】")
    add_body(doc, "")
    add_blank(doc)

    add_heading(doc, "七、团队分工（初步）")
    add_info_table(doc, [
        ("项目经理/负责人", ""),
        ("前端开发", ""),
        ("后端开发", ""),
        ("数据库/架构", ""),
        ("UI/UX 设计", ""),
        ("测试/文档", ""),
    ])

    add_heading(doc, "八、待办事项（Action Items）")
    add_action_table(doc)

    add_heading(doc, "九、风险与问题")
    add_risk_table(doc)

    add_heading(doc, "十、需上报/需协调事项")
    add_body(doc, "1. ")
    add_body(doc, "2. ")
    add_blank(doc)

    add_heading(doc, "十一、下次会议安排")
    add_info_table(doc, [
        ("会议时间", ""),
        ("会议地点/方式", ""),
        ("主要议题", "需求分析 / 原型设计 / 技术方案评审 等"),
        ("需提前准备材料", ""),
    ])

    add_heading(doc, "十二、附件")
    add_body(doc, "□ 头脑风暴便签/白板截图")
    add_body(doc, "□ 参考项目链接或竞品分析")
    add_body(doc, "□ 其他：________")
    add_blank(doc, 2)

    # 签字栏
    p = doc.add_paragraph()
    run = p.add_run("记录人签字：________________    主持人确认：________________    分发日期：YYYY-MM-DD")
    set_run_font(run, FONT_SIZE_SUB)
    add_blank(doc)

    p = doc.add_paragraph()
    run = p.add_run("— 填写说明：请在会后 24 小时内完成记录分发；待办事项须明确负责人与截止日期。 —")
    set_run_font(run, FONT_SIZE_NOTE, color=RGBColor(128, 128, 128))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"已生成：{OUTPUT}")


if __name__ == "__main__":
    main()
