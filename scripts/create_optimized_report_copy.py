"""Create a NEW desktop docx with optimized diagrams. Never overwrite the original."""

from __future__ import annotations

import os
import shutil

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Rectangle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

SRC = r"c:\Users\17879\Desktop\8002124134-酒店预订管理系统-大作业报告.docx"
DST = r"c:\Users\17879\Desktop\8002124134-酒店预订管理系统-大作业报告-图表优化版.docx"
DIAGRAM_DIR = r"c:\Users\17879\Desktop\report_diagrams_v2"

plt.rcParams["font.sans-serif"] = ["SimSun", "Microsoft YaHei", "SimHei"]
plt.rcParams["axes.unicode_minus"] = False


def save(name: str) -> str:
    os.makedirs(DIAGRAM_DIR, exist_ok=True)
    path = os.path.join(DIAGRAM_DIR, name)
    plt.savefig(path, dpi=220, bbox_inches="tight", facecolor="white", pad_inches=0.2)
    plt.close()
    return path


def rounded(ax, x, y, w, h, text, fc="#E8F4FD", ec="#222", fs=9, bold=False):
    ax.add_patch(
        FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.02",
            linewidth=1.1,
            edgecolor=ec,
            facecolor=fc,
            zorder=3,
        )
    )
    ax.text(
        x + w / 2,
        y + h / 2,
        text,
        ha="center",
        va="center",
        fontsize=fs,
        fontweight="bold" if bold else "normal",
        zorder=4,
    )


def arrow(ax, x1, y1, x2, y2):
    ax.add_patch(
        FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="-|>",
            mutation_scale=11,
            lw=1.15,
            color="#222",
            shrinkA=0,
            shrinkB=0,
            zorder=2,
        )
    )


def hline(ax, x1, x2, y):
    ax.plot([x1, x2], [y, y], color="#222", lw=1.15, zorder=2)


def vline(ax, x, y1, y2):
    ax.plot([x, x], [y1, y2], color="#222", lw=1.15, zorder=2)


def tag(ax, x, y, text):
    ax.text(
        x,
        y,
        text,
        ha="center",
        va="center",
        fontsize=8,
        zorder=5,
        bbox=dict(boxstyle="round,pad=0.18", fc="#FFFFFF", ec="#AAAAAA", lw=0.6),
    )


# ---------- 图1 功能结构图 ----------
def fig1():
    fig, ax = plt.subplots(figsize=(11, 6.4))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 6.4)
    ax.axis("off")

    rounded(ax, 4.0, 5.55, 3.0, 0.55, "酒店预订管理系统", fc="#D6EAF8", fs=12, bold=True)

    cols = [
        (0.4, "用户端", ["首页搜索", "酒店列表/详情", "预订下单", "订单中心", "收藏/优惠券", "评价/个人中心"]),
        (4.0, "管理端", ["仪表盘", "酒店审核", "订单管理", "用户管理", "评价/Banner", "优惠券/日志"]),
        (7.6, "商家端", ["经营看板", "我的酒店", "房型管理", "库存日历", "订单处理", "评价回复"]),
    ]
    centers = []
    for x, title, items in cols:
        rounded(ax, x, 4.65, 2.8, 0.48, title, fc="#FCF3CF", fs=10, bold=True)
        for i, item in enumerate(items):
            rounded(ax, x, 4.0 - i * 0.55, 2.8, 0.45, item, fs=8.5)
        centers.append(x + 1.4)

    # 系统框底边中心 -> 横向总线 -> 三端子模块
    bus_y = 5.35
    vline(ax, 5.5, 5.55, bus_y)
    hline(ax, centers[0], centers[2], bus_y)
    for cx in centers:
        arrow(ax, cx, bus_y, cx, 5.13)

    ax.text(0.2, 0.15, "图1 项目功能结构图", fontsize=9, color="#555")
    return save("fig1.png")


# ---------- 图2 业务流程图 ----------
def fig2():
    fig, ax = plt.subplots(figsize=(12.5, 3.2))
    ax.set_xlim(0, 12.5)
    ax.set_ylim(0, 3.2)
    ax.axis("off")

    steps = [
        "①搜索酒店",
        "②浏览详情\n选择房型",
        "③填写入住人\n提交订单",
        "④模拟支付",
        "⑤商家审核\n入住人",
        "⑥入住完成\n退订处理",
        "⑦发表评价",
    ]
    w, h, y = 1.4, 1.05, 1.0
    gap = 0.25
    xs = [0.35 + i * (w + gap) for i in range(len(steps))]
    for x, text in zip(xs, steps):
        rounded(ax, x, y, w, h, text, fs=8)
    for i in range(len(steps) - 1):
        arrow(ax, xs[i] + w, y + h / 2, xs[i + 1], y + h / 2)

    ax.text(0.2, 0.2, "图2 主干业务流程图", fontsize=9, color="#555")
    return save("fig2.png")


# ---------- 图3 用例图 ----------
def fig3():
    fig, ax = plt.subplots(figsize=(12, 8.2))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8.2)
    ax.axis("off")

    # 系统边界
    ax.add_patch(Rectangle((2.9, 0.55), 8.4, 7.0, fill=False, lw=1.6, edgecolor="#222", zorder=1))
    ax.text(7.1, 7.75, "酒店预订管理系统", ha="center", fontsize=12, fontweight="bold")

    # 参与者
    actors = [
        ("注册用户", 6.2),
        ("酒店商家", 3.9),
        ("平台管理员", 1.4),
    ]
    aw, ah = 1.7, 0.55
    for name, ay in actors:
        rounded(ax, 0.35, ay, aw, ah, name, fc="#FADBD8", fs=9, bold=True)

    uw, uh = 2.35, 0.5
    user_cases = [
        (3.3, 6.7, "酒店搜索与浏览"),
        (3.3, 5.95, "预订下单"),
        (3.3, 5.2, "订单支付/取消"),
        (3.3, 4.45, "退订申请"),
        (3.3, 3.7, "评价/收藏/领券"),
    ]
    merchant_cases = [
        (8.5, 6.4, "房型与库存管理"),
        (8.5, 5.5, "订单审核处理"),
        (8.5, 4.6, "退订审核处理"),
        (8.5, 3.7, "评价回复"),
    ]
    admin_cases = [
        (3.3, 1.35, "酒店审核/营销管理"),
        (8.5, 1.35, "用户/订单监管"),
    ]
    for x, y, t in user_cases + merchant_cases + admin_cases:
        rounded(ax, x, y, uw, uh, t, fs=8)

    actor_x = 0.35 + aw
    # 用户：左侧竖总线 -> 各用例
    bus_u = 2.75
    uys = [y + uh / 2 for _, y, _ in user_cases]
    hline(ax, actor_x, bus_u, 6.2 + ah / 2)
    vline(ax, bus_u, min(uys), max(uys))
    for cy in uys:
        arrow(ax, bus_u, cy, 3.3, cy)

    # 商家：绕开左侧用例，从下方通道到右侧总线
    bus_m = 8.15
    my = 3.9 + ah / 2
    route_y = 3.15
    hline(ax, actor_x, 2.4, my)
    vline(ax, 2.4, my, route_y)
    hline(ax, 2.4, bus_m, route_y)
    mys = [y + uh / 2 for _, y, _ in merchant_cases]
    vline(ax, bus_m, route_y, max(mys))
    for cy in mys:
        arrow(ax, bus_m, cy, 8.5, cy)

    # 管理员：分别连左右两个用例，走底部通道
    ay = 1.4 + ah / 2
    # 左用例
    hline(ax, actor_x, bus_u, ay)
    arrow(ax, bus_u, ay, 3.3, 1.35 + uh / 2)
    # 右用例：底部绕行，不穿过左侧用例
    bottom = 0.95
    vline(ax, bus_u, ay, bottom)
    hline(ax, bus_u, bus_m, bottom)
    vline(ax, bus_m, bottom, 1.35 + uh / 2)
    arrow(ax, bus_m, 1.35 + uh / 2, 8.5, 1.35 + uh / 2)

    ax.text(0.2, 0.15, "图3 系统用例图", fontsize=9, color="#555")
    return save("fig3.png")


# ---------- 图4 顶层DFD ----------
def fig4():
    fig, ax = plt.subplots(figsize=(10.5, 5.8))
    ax.set_xlim(0, 10.5)
    ax.set_ylim(0, 5.8)
    ax.axis("off")

    sx, sy, sw, sh = 4.2, 1.5, 3.2, 2.6
    rounded(ax, sx, sy, sw, sh, "0\n酒店预订管理系统", fc="#D5F5E3", fs=12, bold=True)

    ew, eh = 1.8, 0.65
    entities = [
        ("注册用户", 4.2),
        ("酒店商家", 2.55),
        ("平台管理员", 0.9),
    ]
    for name, ey in entities:
        rounded(ax, 0.4, ey, ew, eh, name, fc="#FCF3CF", fs=9)

    # 用户 -> 系统（正交：先水平再向下进入系统框上半部）
    y_user = 4.2 + eh / 2
    y_sys_top = sy + sh - 0.35
    hline(ax, 0.4 + ew, sx - 0.15, y_user)
    vline(ax, sx - 0.15, y_user, y_sys_top)
    arrow(ax, sx - 0.15, y_sys_top, sx, y_sys_top)
    tag(ax, 2.5, y_user + 0.28, "搜索/预订/支付")

    # 商家 <-> 系统（两条水平线，均落在系统框高度内）
    y2a = 2.55 + eh / 2 + 0.15
    y2b = 2.55 + eh / 2 - 0.15
    arrow(ax, 0.4 + ew, y2a, sx, y2a)
    tag(ax, 2.5, y2a + 0.28, "房型/库存/订单")
    arrow(ax, sx, y2b, 0.4 + ew, y2b)
    tag(ax, 2.5, y2b - 0.28, "订单/酒店信息")

    # 管理员 -> 系统（正交：先水平再向上进入系统框下半部）
    y_admin = 0.9 + eh / 2
    y_sys_bot = sy + 0.35
    hline(ax, 0.4 + ew, sx - 0.15, y_admin)
    vline(ax, sx - 0.15, y_admin, y_sys_bot)
    arrow(ax, sx - 0.15, y_sys_bot, sx, y_sys_bot)
    tag(ax, 2.5, y_admin + 0.28, "审核/营销配置")

    ax.text(0.2, 0.15, "图4 顶层数据流图", fontsize=9, color="#555")
    return save("fig4.png")


# ---------- 图5 预订下单 ----------
def fig5():
    fig, ax = plt.subplots(figsize=(11, 5.8))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 5.8)
    ax.axis("off")

    rounded(ax, 0.4, 2.7, 1.2, 0.65, "用户", fc="#FCF3CF")
    rounded(ax, 2.2, 2.5, 2.1, 1.05, "1.0 预订处理", fc="#D6EAF8", fs=10, bold=True)

    stores = [
        (6.2, 4.2, "D1 用户表"),
        (6.2, 2.85, "D2 库存表"),
        (6.2, 1.5, "D3 订单表"),
        (2.35, 0.55, "D4 优惠券表"),
    ]
    for x, y, t in stores:
        rounded(ax, x, y, 1.7, 0.8, t, fc="#E8DAEF", fs=8.5)

    # 用户 -> 处理
    arrow(ax, 1.6, 3.025, 2.2, 3.025)
    tag(ax, 1.9, 3.35, "预订请求")

    # 处理 -> 右侧三表：先到主干，再分支
    trunk = 5.5
    py = 3.025
    hline(ax, 4.3, trunk, py)
    store_ys = [4.6, 3.25, 1.9]
    labels = ["读取用户", "校验库存", "写入订单"]
    vline(ax, trunk, min(store_ys), max(store_ys))
    for sy, lbl in zip(store_ys, labels):
        arrow(ax, trunk, sy, 6.2, sy)
        tag(ax, 5.85, sy + 0.28, lbl)

    # 处理 -> 优惠券
    arrow(ax, 3.25, 2.5, 3.25, 1.35)
    tag(ax, 4.05, 1.95, "抵扣计算")

    ax.text(0.2, 0.15, "图5 预订下单二层数据流图", fontsize=9, color="#555")
    return save("fig5.png")


# ---------- 图6 支付流程 ----------
def fig6():
    fig, ax = plt.subplots(figsize=(11, 5.2))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 5.2)
    ax.axis("off")

    # 实体坐标（左、上、宽、高）
    user = (0.4, 2.15, 1.3, 0.7)
    proc = (2.4, 2.0, 2.3, 1.0)
    sim = (5.4, 3.55, 2.0, 0.7)
    d3 = (8.0, 2.85, 1.8, 0.8)
    d4 = (8.0, 1.35, 1.8, 0.8)

    rounded(ax, *user, "用户", fc="#FCF3CF", fs=9)
    rounded(ax, *proc, "2.0 支付处理", fc="#D6EAF8", fs=10, bold=True)
    rounded(ax, *sim, "模拟支付成功", fc="#FFF8E1", fs=9, bold=True)
    rounded(ax, *d3, "D3 订单表", fc="#E8DAEF", fs=8.5)
    rounded(ax, *d4, "D4 优惠券表", fc="#E8DAEF", fs=8.5)

    # 用户右缘中心 -> 处理左缘中心
    ux, uy, uw, uh = user
    px, py, pw, ph = proc
    sx, sy, sw, sh = sim
    d3x, d3y, d3w, d3h = d3
    d4x, d4y, d4w, d4h = d4

    y_mid = uy + uh / 2
    arrow(ax, ux + uw, y_mid, px, y_mid)
    tag(ax, (ux + uw + px) / 2, y_mid + 0.32, "支付请求")

    # 处理上缘中心 -> 模拟支付成功下缘中心
    proc_top_x = px + pw / 2
    proc_top_y = py + ph
    sim_bot_x = sx + sw / 2
    sim_bot_y = sy
    # 先上再右再上，保证贴边
    hline(ax, proc_top_x, proc_top_x, proc_top_y)  # noop guard
    vline(ax, proc_top_x, proc_top_y, (proc_top_y + sim_bot_y) / 2)
    hline(ax, proc_top_x, sim_bot_x, (proc_top_y + sim_bot_y) / 2)
    arrow(ax, sim_bot_x, (proc_top_y + sim_bot_y) / 2, sim_bot_x, sim_bot_y)
    tag(ax, (proc_top_x + sim_bot_x) / 2 + 0.15, (proc_top_y + sim_bot_y) / 2 + 0.28, "调用模拟支付")

    # 模拟支付成功右缘 -> 主干 -> D3 / D4
    trunk_x = 7.5
    sim_right_x = sx + sw
    sim_mid_y = sy + sh / 2
    hline(ax, sim_right_x, trunk_x, sim_mid_y)
    d3_mid = d3y + d3h / 2
    d4_mid = d4y + d4h / 2
    vline(ax, trunk_x, d4_mid, sim_mid_y)
    arrow(ax, trunk_x, d3_mid, d3x, d3_mid)
    tag(ax, (trunk_x + d3x) / 2, d3_mid + 0.28, "更新为PAID")
    arrow(ax, trunk_x, d4_mid, d4x, d4_mid)
    tag(ax, (trunk_x + d4x) / 2, d4_mid + 0.28, "标记已使用")

    ax.text(0.2, 0.15, "图6 支付流程二层数据流图", fontsize=9, color="#555")
    return save("fig6.png")


# ---------- 图7 订单管理 ----------
def fig7():
    fig, ax = plt.subplots(figsize=(11, 5.6))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 5.6)
    ax.axis("off")

    rounded(ax, 0.4, 4.0, 1.2, 0.65, "用户", fc="#FCF3CF")
    rounded(ax, 0.4, 1.1, 1.2, 0.65, "商家", fc="#FCF3CF")
    rounded(ax, 2.3, 2.35, 2.2, 1.0, "3.0 订单管理", fc="#D6EAF8", fs=10, bold=True)

    rounded(ax, 6.0, 3.6, 1.7, 0.8, "D3 订单表", fc="#E8DAEF", fs=8.5)
    rounded(ax, 6.0, 2.35, 1.7, 0.8, "D2 库存表", fc="#E8DAEF", fs=8.5)
    rounded(ax, 6.0, 1.1, 1.7, 0.8, "D5 日志表", fc="#E8DAEF", fs=8.5)

    # 用户 -> 处理（正交）
    hline(ax, 1.6, 2.0, 4.325)
    vline(ax, 2.0, 4.325, 3.35)
    arrow(ax, 2.0, 3.35, 2.3, 3.35)
    tag(ax, 1.7, 3.9, "取消/退订")

    # 商家 -> 处理（正交）
    hline(ax, 1.6, 2.0, 1.425)
    vline(ax, 2.0, 1.425, 2.55)
    arrow(ax, 2.0, 2.55, 2.3, 2.55)
    tag(ax, 1.7, 1.9, "审核入住/退订")

    # 处理 -> 三表（主干分支）
    trunk = 5.4
    hline(ax, 4.5, trunk, 2.85)
    store_ys = [4.0, 2.75, 1.5]
    labels = ["更新状态", "释放库存", "记录操作"]
    vline(ax, trunk, min(store_ys), max(store_ys))
    for sy, lbl in zip(store_ys, labels):
        arrow(ax, trunk, sy, 6.0, sy)
        tag(ax, 5.7, sy + 0.28, lbl)

    ax.text(0.2, 0.15, "图7 订单管理二层数据流图", fontsize=9, color="#555")
    return save("fig7.png")


def replace_in_copy():
    if not os.path.exists(SRC):
        raise FileNotFoundError(SRC)

    shutil.copy2(SRC, DST)
    print("copied ->", DST)

    images = {
        47: fig1(),
        69: fig2(),
        76: fig3(),
        81: fig4(),
        86: fig5(),
        91: fig6(),
        96: fig7(),
    }

    doc = Document(DST)
    for idx, path in images.items():
        p = doc.paragraphs[idx]
        # 清空原有 runs（含旧图）
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
        run = p.add_run()
        run.add_picture(path, width=Inches(6.3))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("replaced image at paragraph", idx, "->", os.path.basename(path))

    doc.save(DST)
    print("done:", DST)
    print("original untouched:", SRC)


if __name__ == "__main__":
    replace_in_copy()
