"""Update hotel reservation system report docx - sections 1-2 only."""

from __future__ import annotations

import os
import shutil
from io import BytesIO

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = r"c:\Users\17879\Desktop\8002124134-酒店预订管理系统-大作业报告 (已自动恢复).docx"
OUTPUT_PATH = r"c:\Users\17879\Desktop\8002124134-酒店预订管理系统-大作业报告.docx"
BACKUP_PATH = DOC_PATH.replace(".docx", "_backup.docx")
DIAGRAM_DIR = r"c:\Users\17879\Desktop\report_diagrams"

plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False


def save_fig(name: str):
    path = os.path.join(DIAGRAM_DIR, name)
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def box(ax, x, y, w, h, text, fc="#E8F4FD", ec="#333333", fontsize=9):
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.02",
        linewidth=1.2,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, wrap=True)


def arrow(ax, x1, y1, x2, y2):
    ax.add_patch(
        FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.2,
            color="#333333",
            shrinkA=2,
            shrinkB=2,
        )
    )


def replace_paragraph_image(paragraph, image_path, width=Inches(6.2)):
    for run in paragraph.runs:
        run.clear()
    if not paragraph.runs:
        run = paragraph.add_run()
    else:
        run = paragraph.runs[0]
    run.add_picture(image_path, width=width)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_paragraph_text(paragraph, text, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = bold


def gen_function_structure():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    box(ax, 3.5, 5.1, 3, 0.6, "酒店预订管理系统", fc="#D6EAF8", fontsize=11)
    for i, (x, title) in enumerate([(0.2, "用户端"), (3.5, "管理端"), (6.8, "商家端")]):
        box(ax, x, 4.0, 2.5, 0.55, title, fc="#FCF3CF")
    user_items = ["首页搜索", "酒店列表/详情", "预订下单", "订单中心", "收藏/优惠券", "评价/个人中心"]
    admin_items = ["仪表盘", "酒店审核", "订单管理", "用户管理", "评价/Banner", "优惠券/日志"]
    merchant_items = ["经营看板", "我的酒店", "房型管理", "库存日历", "订单处理", "评价回复"]
    for idx, item in enumerate(user_items):
        box(ax, 0.2, 3.2 - idx * 0.45, 2.5, 0.38, item, fontsize=8)
    for idx, item in enumerate(admin_items):
        box(ax, 3.5, 3.2 - idx * 0.45, 2.5, 0.38, item, fontsize=8)
    for idx, item in enumerate(merchant_items):
        box(ax, 6.8, 3.2 - idx * 0.45, 2.5, 0.38, item, fontsize=8)
    arrow(ax, 5, 5.1, 1.45, 4.55)
    arrow(ax, 5, 5.1, 4.75, 4.55)
    arrow(ax, 5, 5.1, 8.05, 4.55)
    return save_fig("fig1_function_structure.png")


def gen_business_flow():
    fig, ax = plt.subplots(figsize=(11, 2.8))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 3)
    ax.axis("off")
    steps = [
        "搜索酒店",
        "浏览详情\n选择房型",
        "填写入住人\n提交订单",
        "模拟支付",
        "商家审核\n入住人",
        "确认入住\n完成/退订",
        "发表评价",
    ]
    xs = [0.2 + i * 1.5 for i in range(len(steps))]
    for x, text in zip(xs, steps):
        box(ax, x, 1.0, 1.2, 0.9, text, fontsize=8)
    for i in range(len(steps) - 1):
        arrow(ax, xs[i] + 1.2, 1.45, xs[i + 1], 1.45)
    return save_fig("fig2_business_flow.png")


def gen_use_case():
    fig, ax = plt.subplots(figsize=(10, 6.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis("off")
    box(ax, 3.0, 0.8, 4.0, 5.2, "", fc="#FFFFFF")
    ax.text(5.0, 5.7, "酒店预订管理系统", ha="center", fontsize=11, weight="bold")
    cases = [
        (3.4, 5.0, "酒店搜索与浏览"),
        (3.4, 4.3, "预订下单"),
        (3.4, 3.6, "订单支付/取消"),
        (3.4, 2.9, "退订申请"),
        (3.4, 2.2, "评价/收藏/领券"),
        (3.4, 1.5, "酒店审核/营销管理"),
        (6.0, 5.0, "房型与库存管理"),
        (6.0, 4.3, "订单审核处理"),
        (6.0, 3.6, "评价回复"),
        (6.0, 2.9, "用户/订单监管"),
    ]
    for x, y, text in cases:
        box(ax, x, y, 2.0, 0.5, text, fontsize=8)
    actors = [(0.5, 4.8, "注册用户"), (0.5, 2.0, "酒店商家"), (0.5, 0.8, "平台管理员")]
    for x, y, text in actors:
        box(ax, x, y, 1.5, 0.55, text, fc="#FADBD8", fontsize=9)
        if text == "注册用户":
            arrow(ax, 2.0, 5.05, 3.4, 5.25)
            arrow(ax, 2.0, 5.05, 3.4, 4.55)
            arrow(ax, 2.0, 4.85, 3.4, 3.85)
        elif text == "酒店商家":
            arrow(ax, 2.0, 2.25, 6.0, 5.25)
            arrow(ax, 2.0, 2.25, 6.0, 4.55)
            arrow(ax, 2.0, 2.25, 6.0, 3.85)
        else:
            arrow(ax, 2.0, 1.05, 3.4, 1.75)
            arrow(ax, 2.0, 1.05, 6.0, 3.15)
    return save_fig("fig3_use_case.png")


def gen_dfd_context():
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 4.5)
    ax.axis("off")
    box(ax, 3.0, 1.5, 3.0, 1.5, "酒店预订\n管理系统", fc="#D5F5E3", fontsize=11)
    box(ax, 0.2, 3.0, 1.6, 0.7, "注册用户", fc="#FCF3CF")
    box(ax, 0.2, 1.8, 1.6, 0.7, "酒店商家", fc="#FCF3CF")
    box(ax, 0.2, 0.6, 1.6, 0.7, "平台管理员", fc="#FCF3CF")
    arrow(ax, 1.8, 3.35, 3.0, 2.7)
    ax.text(2.2, 3.5, "搜索/预订/支付", fontsize=8)
    arrow(ax, 3.0, 2.2, 1.8, 2.15)
    ax.text(2.0, 2.35, "订单/酒店信息", fontsize=8)
    arrow(ax, 1.8, 2.15, 3.0, 2.0)
    ax.text(2.0, 1.95, "房型/库存/订单", fontsize=8)
    arrow(ax, 1.8, 0.95, 3.0, 1.8)
    ax.text(2.0, 1.25, "审核/配置", fontsize=8)
    return save_fig("fig4_dfd_context.png")


def gen_dfd_booking():
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.5)
    ax.axis("off")
    box(ax, 0.2, 2.0, 1.2, 0.7, "用户", fc="#FCF3CF")
    box(ax, 2.0, 2.0, 1.8, 0.8, "1.0\n预订处理", fc="#D6EAF8")
    for i, (x, t) in enumerate([(4.5, "D1 用户"), (6.3, "D2 库存"), (8.1, "D3 订单")]):
        box(ax, x, 1.8, 1.4, 1.0, t, fc="#E8DAEF")
    arrow(ax, 1.4, 2.35, 2.0, 2.35)
    ax.text(1.5, 2.5, "预订请求", fontsize=8)
    arrow(ax, 3.8, 2.45, 4.5, 2.3)
    arrow(ax, 3.8, 2.25, 6.3, 2.3)
    ax.text(4.0, 2.55, "校验/扣减", fontsize=8)
    arrow(ax, 3.8, 2.05, 8.1, 2.3)
    ax.text(5.8, 2.05, "写入订单", fontsize=8)
    arrow(ax, 2.9, 2.0, 2.9, 1.2)
    box(ax, 2.2, 0.5, 1.4, 0.6, "D4 优惠券", fc="#E8DAEF")
    ax.text(3.5, 1.55, "抵扣计算", fontsize=8)
    return save_fig("fig5_dfd_booking.png")


def gen_dfd_payment():
    fig, ax = plt.subplots(figsize=(9, 4))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 4)
    ax.axis("off")
    box(ax, 0.2, 2.0, 1.2, 0.7, "用户", fc="#FCF3CF")
    box(ax, 2.0, 2.0, 1.8, 0.8, "2.0\n支付处理", fc="#D6EAF8")
    box(ax, 4.8, 2.0, 1.5, 0.9, "D3 订单", fc="#E8DAEF")
    box(ax, 6.8, 2.0, 1.5, 0.9, "D4 优惠券", fc="#E8DAEF")
    arrow(ax, 1.4, 2.35, 2.0, 2.35)
    ax.text(1.5, 2.5, "支付请求", fontsize=8)
    arrow(ax, 3.8, 2.45, 4.8, 2.45)
    ax.text(4.0, 2.6, "更新状态", fontsize=8)
    arrow(ax, 3.8, 2.15, 6.8, 2.35)
    ax.text(5.2, 2.0, "标记已使用", fontsize=8)
    arrow(ax, 2.9, 2.8, 2.9, 3.3)
    ax.text(2.2, 3.15, "模拟支付成功\n(PAID)", fontsize=8)
    return save_fig("fig6_dfd_payment.png")


def gen_dfd_order():
    fig, ax = plt.subplots(figsize=(10, 4.8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.8)
    ax.axis("off")
    box(ax, 0.2, 3.0, 1.2, 0.7, "用户", fc="#FCF3CF")
    box(ax, 0.2, 1.2, 1.2, 0.7, "商家", fc="#FCF3CF")
    box(ax, 2.0, 2.0, 2.0, 0.9, "3.0\n订单管理", fc="#D6EAF8")
    box(ax, 4.8, 2.0, 1.5, 0.9, "D3 订单", fc="#E8DAEF")
    box(ax, 6.8, 2.0, 1.5, 0.9, "D2 库存", fc="#E8DAEF")
    box(ax, 8.3, 0.8, 1.4, 0.9, "D5 日志", fc="#E8DAEF")
    arrow(ax, 1.4, 3.35, 2.0, 2.6)
    ax.text(1.5, 3.5, "取消/退订", fontsize=8)
    arrow(ax, 1.4, 1.55, 2.0, 2.3)
    ax.text(1.5, 1.9, "审核入住/退订", fontsize=8)
    arrow(ax, 4.0, 2.45, 4.8, 2.45)
    arrow(ax, 4.0, 2.15, 6.8, 2.35)
    arrow(ax, 4.0, 1.95, 8.3, 1.25)
    ax.text(5.0, 1.7, "释放库存/记录操作", fontsize=8)
    return save_fig("fig7_dfd_order.png")


def update_document():
    os.makedirs(DIAGRAM_DIR, exist_ok=True)
    if not os.path.exists(BACKUP_PATH):
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    diagrams = {
        47: gen_function_structure(),
        69: gen_business_flow(),
        76: gen_use_case(),
        81: gen_dfd_context(),
        86: gen_dfd_booking(),
        91: gen_dfd_payment(),
        96: gen_dfd_order(),
    }

    doc = Document(DOC_PATH)
    paras = doc.paragraphs

    text_updates = {
        41: "随着旅游业的快速发展和人们出行需求的日益增长，酒店预订已成为日常生活中不可或缺的服务环节。传统的酒店预订方式存在信息不透明、预订效率低、管理成本高等问题。为解决这些痛点，本团队设计并开发了酒店预订管理系统。",
        42: "本系统采用 B/S 架构，基于 Spring Boot 4 构建后端 REST API，使用 MySQL 8 作为关系型数据库、Redis 作为缓存与会话支撑，前端分为用户端（Vue 3 + Vant）和管理端（Vue 3 + Element Plus）。系统面向三类角色：平台管理员负责酒店审核与平台运营；酒店商家负责酒店、房型、库存与订单处理；注册用户完成搜索、预订、支付、评价等全流程操作。数据库结构通过 Flyway 版本化迁移管理，并支持 Docker 一键部署。",
        44: "平台管理端：仪表盘（订单量、GMV、待审核酒店数等统计）、酒店审核（通过/驳回入驻申请）、订单管理（全平台订单查询）、用户管理（用户列表、账号禁用）、评价管理（展示/隐藏评价）、Banner 管理、优惠券管理、操作日志查看。",
        45: "酒店商家端：经营看板（本店订单与酒店统计）、我的酒店（酒店信息 CRUD）、房型管理（房型 CRUD、床型、可住人数）、库存日历（按日设置价格与可售间数）、订单管理（审核入住人信息、处理提前退订申请）、评价管理（查看并回复用户评价）。",
        46: "用户端：首页搜索（省/市、入住离店日期、关键词）、酒店列表（分页、星级筛选、价格排序）、酒店详情（图文、设施、房型、评价摘要）、预订下单（填写入住人、选择优惠券）、订单中心（待支付/已确认/已取消等状态管理）、模拟支付、取消订单、提前退订申请、订单评价、酒店收藏、优惠券领取与使用、个人资料与密码修改。",
        54: "本系统的主要目标是构建一个功能完善、操作便捷的酒店预订管理平台，解决传统酒店预订流程中信息不对称、预订效率低和管理成本高等问题。系统面向平台管理员、酒店商家和注册用户三类角色，提供差异化的功能服务。",
        55: "通过本系统，平台管理员可以高效完成酒店入驻审核和平台数据监管；酒店商家可以灵活维护房型、库存并处理订单；注册用户可以实现从搜索、预订、支付到评价的全流程线上操作，提升酒店行业的数字化管理水平。",
        57: "涉众：本系统的主要操作者包括平台管理员、酒店商家和注册用户（住客）。",
        58: "业务愿景：按涉众分别介绍系统需求，如下表所示。",
        61: "本系统的主要业务流程围绕用户预订酒店展开，核心业务流程如下：",
        62: "（1）用户在首页输入省/市、入住/离店日期和关键词进行搜索，系统展示符合条件的酒店列表，支持按星级筛选和价格排序。",
        63: "（2）用户进入酒店详情页浏览酒店信息、设施服务和用户评价，选择房型后进入预订页面。",
        64: "（3）用户填写入住人信息、选择入住人数和优惠券，系统校验库存与价格后生成待支付订单，并预扣对应日期库存。",
        65: "（4）用户在订单详情页完成模拟支付，订单状态变为“已支付”，等待商家审核入住人信息。",
        66: "（5）商家审核通过后订单变为“已确认”；若拒绝则释放库存并退款。用户可在入住前取消待支付订单，或在已确认状态下申请提前退订。",
        67: "（6）用户完成入住后可对订单发表星级评价，商家可在管理端查看并回复评价。",
        72: "本系统的功能通过用例图来描述，系统的参与者（Actor）分为三类：平台管理员、酒店商家和注册用户（住客）。",
        73: "平台管理员的用例包括：查看运营仪表盘、审核酒店入驻、管理全平台订单、管理用户账号、管理评价展示状态、管理 Banner 与优惠券、查看操作日志等。",
        74: "酒店商家的用例包括：维护本店酒店信息、管理房型、设置日历价格与库存、审核入住人信息、处理退订申请、回复用户评价、查看经营看板等。",
        75: "注册用户（住客）的用例包括：注册登录、搜索浏览酒店、预订下单、模拟支付、管理订单（取消/退订/完成）、发表评价、收藏酒店、领取和使用优惠券、维护个人资料等。",
        79: "本系统采用数据流图（DFD）描述核心业务的数据流程，分为顶层数据流图和二层数据流图。",
        80: "顶层数据流图描述了系统与外部实体（注册用户、酒店商家、平台管理员）之间的数据交互关系。用户向系统提交搜索和预订请求，系统返回酒店信息和订单状态；酒店商家维护房型与库存并处理订单；平台管理员进行审核和营销配置。",
        84: "预订下单的数据流程如下：用户在前端发起预订请求并填写入住人信息。后端校验用户登录状态、房型库存与入住人数，计算价格并应用优惠券抵扣，在订单表中插入状态为“待支付”的记录，同时预扣库存日历中的可售间数。若校验失败则返回错误提示。",
        85: "预订下单涉及的数据存储包括：用户表（sys_user）、库存日历表（inventory_calendar）、订单表（hotel_order）、入住人表（order_guest）和优惠券相关表。",
        89: "支付流程的数据流程如下：用户在订单详情页发起模拟支付请求。后端校验订单状态为“待支付”后，将订单状态更新为“已支付（PAID）”，记录支付时间；若订单使用了优惠券，则同步更新用户优惠券的使用状态。本系统采用模拟支付，不对接真实第三方支付通道。",
        90: "支付流程涉及的数据存储包括：订单表（hotel_order）和用户优惠券表（user_coupon）。",
        94: "订单管理的数据流程如下：用户可取消待支付订单并释放库存；商家审核入住人信息，通过后订单变为“已确认”，拒绝则退款并释放库存；用户可在已确认状态下申请提前退订，商家审核后按退订政策计算退款金额并更新订单状态；用户完成入住后可提交评价。",
        95: "订单管理涉及的数据存储包括：订单表（hotel_order）、库存日历表（inventory_calendar）和操作日志表（sys_audit_log）。",
    }

    toc_updates = {
        13: "1. 选题简介\t1",
        14: "1.1 项目简介\t1",
        15: "1.2 项目功能\t1",
        16: "1.3 项目团队\t1",
        17: "2. 需求分析\t2",
        18: "2.1 系统目标\t2",
        19: "2.2 系统涉众\t2",
        20: "2.3 业务流程\t2",
        21: "2.4 功能定义\t2",
        22: "2.5 数据流程\t2",
        23: "2.5.1 预订下单\t2",
        24: "2.5.2 支付流程\t3",
        25: "2.5.3 订单管理\t3",
        26: "3. 数据库概念模型\t4",
        27: "4. 数据库物理模型\t4",
        28: "5.1 物理模型\t4",
        29: "5.2 表空间创建\t4",
        30: "5.3 用户创建\t4",
        31: "5.4 基础数据生成\t4",
        32: "5. 关键应用编程\t5",
        33: "5.1 会员办理\t5",
        34: "5.2 会员充值\t5",
        35: "5.3 销售买单\t5",
        36: "6. 总结\t5",
    }

    for idx, text in {**text_updates, **toc_updates}.items():
        if idx < len(paras):
            set_paragraph_text(paras[idx], text)

    # Update business vision table - remove unimplemented stakeholders
    table = doc.tables[1]
    while len(table.rows) > 3:
        table._tbl.remove(table.rows[-1]._tr)
    table.rows[2].cells[0].text = "注册用户（住客）"
    table.rows[2].cells[1].text = "搜索浏览酒店、预订下单、模拟支付、管理订单、发表评价、收藏酒店、使用优惠券"
    table.rows[2].cells[2].text = "系统终端使用者"

    for idx, img_path in diagrams.items():
        replace_paragraph_image(paras[idx], img_path)

    doc.save(OUTPUT_PATH)
    print(f"Updated: {OUTPUT_PATH}")
    print(f"Source: {DOC_PATH}")
    print(f"Backup: {BACKUP_PATH}")


if __name__ == "__main__":
    update_document()
