# -*- coding: utf-8 -*-
"""填充会议记录并添加 Word 导航（标题样式 + 目录）"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_OUTPUT = r"c:\Users\17879\Desktop\会议记录.docx"
FONT_NAME = "宋体"


def set_run_font(run, size=Pt(12), bold=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, Pt(11), bold=bold)


def apply_heading_style(paragraph, level):
    """将段落设为 Heading 1/2/3，保留文字并统一中文字体"""
    style_name = f"Heading {level}"
    paragraph.style = style_name
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        if level == 1:
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(12)
            run.font.bold = True


def insert_toc(paragraph):
    """在段落位置插入 Word 自动目录域"""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "右键单击此处 → 更新域，即可生成目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_text)
    r.append(fld_end)


def replace_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def fill_bullet_paragraph(paragraph, text):
    replace_paragraph_text(paragraph, text)
    for run in paragraph.runs:
        set_run_font(run, Pt(12))


def find_theme_paragraph(doc, keyword):
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text:
            return i
    return None


def has_toc(doc):
    return any(p.text.strip() in ("目  录", "目录") for p in doc.paragraphs[:6])


def main():
    doc = Document(INPUT_OUTPUT)

    # ── 1. 设置标题样式（启用导航窗格） ──
    heading1_keywords = ("一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、", "十一、", "十二、")
    heading2_keywords = ("4.1", "4.2", "4.3", "5.1", "5.2")

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text.startswith(heading1_keywords):
            apply_heading_style(p, 1)
        elif text.startswith(heading2_keywords):
            apply_heading_style(p, 2)

    # 文档主标题设为 Title
    title_p = doc.paragraphs[0]
    title_p.style = doc.styles["Title"]
    for run in title_p.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        run.font.size = Pt(16)
        run.font.bold = True
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 2. 在标题后插入目录（仅首次） ──
    if not has_toc(doc):
        toc_heading = doc.paragraphs[1].insert_paragraph_before()
        toc_heading.style = doc.styles["Heading 1"]
        run = toc_heading.add_run("目  录")
        set_run_font(run, Pt(14), bold=True)
        toc_heading.paragraph_format.space_after = Pt(6)

        toc_field = doc.paragraphs[2].insert_paragraph_before()
        insert_toc(toc_field)
        toc_field.paragraph_format.space_after = Pt(12)

        page_break = doc.paragraphs[3].insert_paragraph_before()
        run = page_break.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    # ── 3. 填充表格 ──
    # Table 0: 基本信息
    t0 = doc.tables[0]
    rows_data_0 = [
        ("会议主题", "全栈项目主题头脑风暴与选型讨论"),
        ("会议类型", "项目立项"),
        ("会议时间", "2026-07-07  14:00 — 18:00"),
        ("会议地点/方式", "线下（软件楼227）"),
        ("参会人员", "郝帅、何嘉明、方泉顺、谢宇韬、冯硕成"),
        ("主持人 / 记录人", "方泉顺 / 何嘉明（缺席：无）"),
    ]
    for i, (k, v) in enumerate(rows_data_0):
        set_cell(t0.rows[i].cells[0], k, bold=True)
        set_cell(t0.rows[i].cells[1], v)

    # Table 2: 候选主题
    t2 = doc.tables[2]
    themes = [
        ("1", "外卖点餐系统", "餐厅浏览、购物车、在线下单", "周边白领/学生", "LBS定位、订单实时状态", "中", "2票"),
        ("2", "酒店预订系统", "酒店搜索、房型选择、在线预订", "出差/旅游用户", "日历库存、多角色后台", "高", "4票"),
        ("3", "宿舍管理系统", "报修申请、公告通知、水电查询", "宿管/学生", "工单流转", "低", "1票"),
    ]
    for i, row in enumerate(themes, 1):
        for j, val in enumerate(row):
            set_cell(t2.rows[i].cells[j], val)

    # Table 3: 最终决议
    t3 = doc.tables[3]
    resolution = [
        ("选定项目主题", "酒店预订系统"),
        ("项目一句话描述", "面向出行用户的在线酒店搜索与预订平台"),
        ("目标用户群体", "有差旅、旅游住宿需求的个人用户"),
        ("核心功能范围（MVP）", "注册登录、酒店浏览、日期查询、在线下单、订单管理、后台酒店管理"),
        ("预计开发周期", "4 周"),
        ("未采纳主题及原因", "外卖：竞品多、支付对接复杂；宿舍：场景窄、扩展性不足"),
    ]
    for i, (k, v) in enumerate(resolution):
        set_cell(t3.rows[i].cells[0], k, bold=True)
        set_cell(t3.rows[i].cells[1], v)

    # Table 4: 技术栈
    t4 = doc.tables[4]
    tech = [
        ("前端", "Vue / React / Angular 等", "Vue 3 + Element Plus", "组件化开发"),
        ("后端", "Spring Boot / Node.js 等", "Spring Boot", "RESTful API"),
        ("数据库", "MySQL / PostgreSQL 等", "MySQL", "主数据存储"),
        ("缓存/中间件", "Redis / RabbitMQ 等", "Redis", "会话与热点缓存"),
        ("部署", "Docker / 云服务器 等", "Docker + 云服务器", "容器化部署"),
    ]
    for i, row in enumerate(tech, 1):
        for j, val in enumerate(row):
            set_cell(t4.rows[i].cells[j], val, bold=(j == 0))

    # Table 7: 下次会议 — 精简议题描述
    t7 = doc.tables[7]
    set_cell(t7.rows[2].cells[1], "需求分析与数据库设计评审")
    set_cell(t7.rows[3].cells[1], "功能模块清单、初版需求文档")

    # ── 4. 填充段落 ──
    # 背景说明精简
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "为了完成团队激励任务，我们组展开激烈的讨论，想要想出一个牛逼的项目":
            replace_paragraph_text(p, "为完成团队激励任务，小组讨论并筛选全栈练手项目，兼顾业务完整性与开发可行性。")
            fill_bullet_paragraph(p, p.text)
        elif t.startswith("【主题一："):
            replace_paragraph_text(p, "【主题一：外卖点餐系统】")
        elif t.startswith("【主题二："):
            replace_paragraph_text(p, "【主题二：酒店预订系统】")
        elif t.startswith("【主题三："):
            replace_paragraph_text(p, "【主题三：宿舍管理系统】")

    # 按顺序填充三个主题的讨论要点（找主题块后的空 bullet）
    theme_blocks = [
        {
            "keyword": "主题一",
            "title": "【主题一：外卖点餐系统】",
            "points": [
                "讨论要点：",
                "• 需对接支付与配送，模块多、开发周期长",
                "• 周边同类 App 多，差异化难度大",
                "优势：",
                "• 用户习惯成熟，需求边界清晰",
                "顾虑/风险：",
                "• 同质化严重，全栈亮点不足",
            ],
        },
        {
            "keyword": "主题二",
            "title": "【主题二：酒店预订系统】",
            "points": [
                "讨论要点：",
                "• 覆盖搜索、预订、订单、后台管理，业务链完整",
                "• 可加入优惠券、会员等扩展功能",
                "优势：",
                "• 适合全栈技术展示，团队有参考案例",
                "顾虑/风险：",
                "• 需处理日期库存与防超卖逻辑",
            ],
        },
        {
            "keyword": "主题三",
            "title": "【主题三：宿舍管理系统】",
            "points": [
                "讨论要点：",
                "• 功能偏校内管理，开发量较小",
                "• 业务场景单一，答辩展示亮点有限",
                "优势：",
                "• 开发快，上手成本低",
                "顾虑/风险：",
                "• 扩展性弱，难以体现全栈深度",
            ],
        },
    ]

    for block in theme_blocks:
        idx = find_theme_paragraph(doc, block["keyword"])
        if idx is None:
            continue
        fill_bullet_paragraph(doc.paragraphs[idx], block["title"])
        for offset, content in enumerate(block["points"], 1):
            if idx + offset < len(doc.paragraphs):
                fill_bullet_paragraph(doc.paragraphs[idx + offset], content)

    # 技术选型说明
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "【技术选型说明】" and i + 1 < len(doc.paragraphs):
            fill_bullet_paragraph(
                doc.paragraphs[i + 1],
                "前后端分离架构：Vue 3 负责交互，Spring Boot 提供接口，MySQL 存业务数据，Redis 做缓存。",
            )
            break

    # 需上报/协调事项
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "九、需上报/需协调事项" or "九、需上报" in t:
            # 找后面的 1. 2.
            for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip() == "1.":
                    fill_bullet_paragraph(doc.paragraphs[j], "1. 向指导老师确认酒店预订项目立项")
                elif doc.paragraphs[j].text.strip() == "2.":
                    fill_bullet_paragraph(doc.paragraphs[j], "2. 协调云服务器资源用于部署测试")
            break

    # 精简页脚说明
    for p in doc.paragraphs:
        if "填写说明" in p.text:
            replace_paragraph_text(p, "— 本文档已归档，待办事项见第八节。 —")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, Pt(10), color=RGBColor(128, 128, 128))
            break

    doc.save(INPUT_OUTPUT)
    print(f"已完成：{INPUT_OUTPUT}")
    print("提示：打开 Word 后右键目录 → 更新域，即可显示完整导航。")


if __name__ == "__main__":
    main()
