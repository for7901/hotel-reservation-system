"""Regenerate standardized diagrams and replace images in the report docx."""

from __future__ import annotations

import os

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Rectangle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

OUTPUT_PATH = r"c:\Users\17879\Desktop\8002124134-酒店预订管理系统-大作业报告.docx"
DIAGRAM_DIR = r"c:\Users\17879\Desktop\report_diagrams"

plt.rcParams["font.sans-serif"] = ["SimSun", "Microsoft YaHei", "SimHei"]
plt.rcParams["axes.unicode_minus"] = False
FONT = 9
FONT_TITLE = 11
FONT_LABEL = 8


def save_fig(name: str) -> str:
    path = os.path.join(DIAGRAM_DIR, name)
    plt.savefig(path, dpi=200, bbox_inches="tight", facecolor="white", pad_inches=0.15)
    plt.close()
    return path


def box(ax, x, y, w, h, text, fc="#E8F4FD", ec="#222222", fs=FONT, bold=False):
    patch = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.015,rounding_size=0.015",
        linewidth=1.0, edgecolor=ec, facecolor=fc,
    )
    ax.add_patch(patch)
    ax.text(
        x + w / 2, y + h / 2, text,
        ha="center", va="center", fontsize=fs,
        fontweight="bold" if bold else "normal",
    )
    return (x, y, w, h)


def line(ax, x1, y1, x2, y2):
    ax.plot([x1, x2], [y1, y2], color="#222", lw=1.0, zorder=1)


def arrow_head(ax, x1, y1, x2, y2):
    ax.add_patch(FancyArrowPatch(
        (x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=10,
        linewidth=1.0, color="#222", shrinkA=0, shrinkB=0, zorder=2,
    ))


def label(ax, x, y, text):
    ax.text(x, y, text, fontsize=FONT_LABEL, ha="center", va="center",
            bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="#CCC", lw=0.5))


def connect_h(ax, x1, x2, y, alabel=None, label_x=None):
    line(ax, x1, y, x2, y)
    arrow_head(ax, x2 - 0.01, y, x2, y)
    if alabel:
        label(ax, label_x if label_x else (x1 + x2) / 2, y + 0.18, alabel)


def connect_v(ax, x, y1, y2, alabel=None, label_y=None):
    line(ax, x, y1, x, y2)
    arrow_head(ax, x, y2 - 0.01 if y2 > y1 else y2 + 0.01, x, y2)
    if alabel:
        label(ax, x + 0.25, label_y if label_y else (y1 + y2) / 2, alabel)


def connect_bus(ax, actor_x, bus_x, case_x, case_y, case_h):
    """Actor -> vertical bus -> horizontal stub -> use case."""
    ay = case_y + case_h / 2
    line(ax, actor_x, ay, bus_x, ay)
    line(ax, bus_x, ay, case_x - 0.02, ay)
    arrow_head(ax, case_x - 0.02, ay, case_x, ay)


def gen_function_structure():
    fig, ax = plt.subplots(figsize=(10.5, 6.0))
    ax.set_xlim(0, 10.5)
    ax.set_ylim(0, 6.0)
    ax.axis("off")
    box(ax, 3.75, 5.35, 3.0, 0.55, "酒店预订管理系统", fc="#D6EAF8", fs=FONT_TITLE, bold=True)
    cols = [
        (0.3, "用户端", ["首页搜索", "酒店列表/详情", "预订下单", "订单中心", "收藏/优惠券", "评价/个人中心"]),
        (3.75, "管理端", ["仪表盘", "酒店审核", "订单管理", "用户管理", "评价/Banner", "优惠券/日志"]),
        (7.2, "商家端", ["经营看板", "我的酒店", "房型管理", "库存日历", "订单处理", "评价回复"]),
    ]
    centers = []
    for x, title, items in cols:
        box(ax, x, 4.55, 2.55, 0.5, title, fc="#FCF3CF", bold=True)
        for i, item in enumerate(items):
            box(ax, x, 4.0 - i * 0.52, 2.55, 0.42, item, fs=8)
        cx = x + 1.275
        centers.append(cx)
        line(ax, cx, 5.35, cx, 4.55 + 0.52)
        arrow_head(ax, cx, 4.55 + 0.52, cx, 4.55 + 0.5)
    line(ax, centers[0], 5.35, centers[2], 5.35)
    line(ax, 5.25, 5.9, 5.25, 5.35)
    arrow_head(ax, 5.25, 5.35, 5.25, 5.35 - 0.01)
    return save_fig("fig1_function_structure.png")


def gen_business_flow():
    fig, ax = plt.subplots(figsize=(12, 3.0))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 3.0)
    ax.axis("off")
    steps = [
        "① 搜索酒店", "② 浏览详情\n选择房型", "③ 填写入住人\n提交订单",
        "④ 模拟支付", "⑤ 商家审核\n入住人", "⑥ 入住完成\n退订处理", "⑦ 发表评价",
    ]
    w, h, y = 1.35, 1.0, 0.9
    xs = [0.3 + i * 1.65 for i in range(len(steps))]
    for x, text in zip(xs, steps):
        box(ax, x, y, w, h, text, fs=8)
    for i in range(len(steps) - 1):
        connect_h(ax, xs[i] + w, xs[i + 1], y + h / 2)
    return save_fig("fig2_business_flow.png")


def gen_use_case():
    fig, ax = plt.subplots(figsize=(11, 7.8))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 7.8)
    ax.axis("off")
    bx, by, bw, bh = 2.6, 0.5, 7.8, 6.8
    ax.add_patch(Rectangle((bx, by), bw, bh, fill=False, edgecolor="#222", lw=1.5))
    ax.text(bx + bw / 2, by + bh + 0.15, "酒店预订管理系统", ha="center", fontsize=FONT_TITLE, fontweight="bold")

    aw, ah = 1.55, 0.55
    actors = [("注册用户", 5.8), ("酒店商家", 3.8), ("平台管理员", 1.6)]
    for name, ay in actors:
        box(ax, 0.35, ay, aw, ah, name, fc="#FADBD8", bold=True)

    uw, uh = 2.3, 0.5
    user_cases = [
        (3.0, 6.3, "酒店搜索与浏览"), (3.0, 5.55, "预订下单"), (3.0, 4.8, "订单支付/取消"),
        (3.0, 4.05, "退订申请"), (3.0, 3.3, "评价/收藏/领券"),
    ]
    merchant_cases = [
        (7.5, 6.0, "房型与库存管理"), (7.5, 5.1, "订单审核处理"),
        (7.5, 4.2, "退订审核处理"), (7.5, 3.3, "评价回复"),
    ]
    admin_cases = [(3.0, 1.2, "酒店审核/营销管理"), (7.5, 1.2, "用户/订单监管")]

    for x, y, t in user_cases + merchant_cases + admin_cases:
        box(ax, x, y, uw, uh, t, fs=8)

    bus_user = 2.55
    actor_x = 0.35 + aw
    user_ys = [y + uh / 2 for _, y, _ in user_cases]
    line(ax, bus_user, min(user_ys), bus_user, max(user_ys))
    for cy in user_ys:
        connect_h(ax, bus_user, 3.0, cy)
    connect_h(ax, actor_x, bus_user, 5.8 + ah / 2)

    bus_r = 7.0
    route_y = 2.85
    my = 3.8 + ah / 2
    line(ax, 0.35 + aw, my, 2.15, my)
    line(ax, 2.15, my, 2.15, route_y)
    line(ax, 2.15, route_y, bus_r, route_y)
    ys = [y + uh / 2 for _, y, _ in merchant_cases]
    line(ax, bus_r, route_y, bus_r, max(ys))
    for cy in ys:
        connect_h(ax, bus_r, 7.5, cy)

    connect_bus(ax, actor_x, bus_user, 3.0, 1.2, uh)
    connect_h(ax, actor_x, bus_user, 1.6 + ah / 2)
    admin_route_y = 1.0
    line(ax, bus_user, 1.6 + ah / 2, bus_user, admin_route_y)
    line(ax, bus_user, admin_route_y, bus_r, admin_route_y)
    connect_h(ax, bus_r, 7.5, 1.45)

    ax.text(0.2, 0.05, "图3 系统用例图", fontsize=FONT_LABEL, color="#666")
    return save_fig("fig3_use_case.png")


def gen_dfd_context():
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.5)
    ax.axis("off")
    sx, sy, sw, sh = 3.8, 1.5, 3.0, 2.2
    box(ax, sx, sy, sw, sh, "0\n酒店预订管理系统", fc="#D5F5E3", fs=FONT_TITLE, bold=True)
    ew, eh = 1.7, 0.65
    entities = [("注册用户", 3.9), ("酒店商家", 2.4), ("平台管理员", 0.9)]
    for name, ey in entities:
        box(ax, 0.35, ey, ew, eh, name, fc="#FCF3CF")

    flows = [
        (2.05, 4.225, 3.8, "搜索/预订/支付"),
        (3.8, 2.725, 2.05, "订单/酒店信息"),
        (2.05, 2.725, 3.8, "房型/库存/订单"),
        (2.05, 1.225, 3.8, "审核/营销配置"),
    ]
    for x1, y1, x2, lbl in flows:
        connect_h(ax, x1, x2, y1, lbl, (x1 + x2) / 2)

    ax.text(0.2, 0.05, "图4 顶层数据流图", fontsize=FONT_LABEL, color="#666")
    return save_fig("fig4_dfd_context.png")


def gen_dfd_booking():
    fig, ax = plt.subplots(figsize=(10.5, 5.5))
    ax.set_xlim(0, 10.5)
    ax.set_ylim(0, 5.5)
    ax.axis("off")
    box(ax, 0.4, 2.5, 1.1, 0.65, "用户", fc="#FCF3CF")
    box(ax, 2.0, 2.35, 2.0, 0.95, "1.0 预订处理", fc="#D6EAF8", bold=True)
    stores = [
        (5.5, 3.8, "D1\n用户表"),
        (5.5, 2.5, "D2\n库存表"),
        (5.5, 1.2, "D3\n订单表"),
        (2.0, 0.5, "D4\n优惠券表"),
    ]
    sw, sh = 1.45, 0.85
    for x, y, t in stores:
        box(ax, x, y, sw, sh, t, fc="#E8DAEF", fs=8)

    py = 2.82
    px = 4.0
    connect_h(ax, 1.5, 2.0, py, "预订请求", 1.75)
    trunk_x = 4.7
    line(ax, px, py, trunk_x, py)
    store_ys = [4.225, 2.925, 1.625]
    labels = ["读取用户", "校验库存", "写入订单"]
    line(ax, trunk_x, min(store_ys), trunk_x, max(store_ys))
    for sy, lbl in zip(store_ys, labels):
        connect_h(ax, trunk_x, 5.5, sy, lbl, (trunk_x + 5.5) / 2)
    connect_v(ax, 3.0, 2.35, 1.35, "抵扣计算", 1.9)
    connect_v(ax, 2.725, 1.35, 0.5 + sh)

    ax.text(0.2, 0.05, "图5 预订下单二层数据流图", fontsize=FONT_LABEL, color="#666")
    return save_fig("fig5_dfd_booking.png")


def gen_dfd_payment():
    fig, ax = plt.subplots(figsize=(9.5, 4.2))
    ax.set_xlim(0, 9.5)
    ax.set_ylim(0, 4.2)
    ax.axis("off")
    box(ax, 0.4, 1.8, 1.1, 0.65, "用户", fc="#FCF3CF")
    box(ax, 2.0, 1.65, 2.0, 0.95, "2.0 支付处理", fc="#D6EAF8", bold=True)
    box(ax, 5.0, 2.1, 1.5, 0.85, "D3\n订单表", fc="#E8DAEF", fs=8)
    box(ax, 5.0, 0.9, 1.5, 0.85, "D4\n优惠券表", fc="#E8DAEF", fs=8)
    py = 2.12
    connect_h(ax, 1.5, 2.0, py, "支付请求", 1.75)
    connect_h(ax, 4.0, 5.0, 2.52, "更新为PAID", 4.5)
    connect_h(ax, 4.0, 5.0, 1.32, "标记已使用", 4.5)
    label(ax, 3.0, 3.2, "模拟支付成功")
    ax.text(0.2, 0.05, "图6 支付流程二层数据流图", fontsize=FONT_LABEL, color="#666")
    return save_fig("fig6_dfd_payment.png")


def gen_dfd_order():
    fig, ax = plt.subplots(figsize=(10.5, 5.5))
    ax.set_xlim(0, 10.5)
    ax.set_ylim(0, 5.5)
    ax.axis("off")
    box(ax, 0.4, 3.8, 1.1, 0.65, "用户", fc="#FCF3CF")
    box(ax, 0.4, 1.0, 1.1, 0.65, "商家", fc="#FCF3CF")
    box(ax, 2.0, 2.2, 2.2, 1.0, "3.0 订单管理", fc="#D6EAF8", bold=True)
    box(ax, 5.2, 3.2, 1.45, 0.85, "D3\n订单表", fc="#E8DAEF", fs=8)
    box(ax, 5.2, 2.0, 1.45, 0.85, "D2\n库存表", fc="#E8DAEF", fs=8)
    box(ax, 5.2, 0.7, 1.45, 0.85, "D5\n日志表", fc="#E8DAEF", fs=8)

    connect_h(ax, 1.5, 2.0, 4.125, "取消/退订", 1.75)
    connect_h(ax, 1.5, 2.0, 1.325, "审核入住/退订", 1.75)
    connect_h(ax, 4.2, 5.2, 3.62, "更新状态", 4.7)
    connect_h(ax, 4.2, 5.2, 2.42, "释放库存", 4.7)
    connect_h(ax, 4.2, 5.2, 1.12, "记录操作", 4.7)

    ax.text(0.2, 0.05, "图7 订单管理二层数据流图", fontsize=FONT_LABEL, color="#666")
    return save_fig("fig7_dfd_order.png")


def replace_images():
    os.makedirs(DIAGRAM_DIR, exist_ok=True)
    generators = [
        gen_function_structure, gen_business_flow, gen_use_case,
        gen_dfd_context, gen_dfd_booking, gen_dfd_payment, gen_dfd_order,
    ]
    indices = [47, 69, 76, 81, 86, 91, 96]
    doc = Document(OUTPUT_PATH)
    for idx, gen in zip(indices, generators):
        path = gen()
        p = doc.paragraphs[idx]
        for run in p.runs:
            run.clear()
        run = p.add_run()
        run.add_picture(path, width=Inches(6.3))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    targets = [OUTPUT_PATH, OUTPUT_PATH.replace(".docx", "_图表优化.docx")]
    saved = False
    for target in targets:
        try:
            doc.save(target)
            print("saved", target)
            saved = True
            break
        except PermissionError:
            continue
    if not saved:
        alt = OUTPUT_PATH.replace(".docx", "_图表优化2.docx")
        doc.save(alt)
        print("saved", alt)


if __name__ == "__main__":
    replace_images()
